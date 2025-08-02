# === Standard Library Imports ===
import os
import re
import time
import tkinter as tk
from tkinter import filedialog, messagebox

# === Third-Party Library Imports ===
import fitz
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from ollama import chat
from ollama import ChatResponse
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from tkinterdnd2 import TkinterDnD, DND_FILES

# === Global Prompt Storage ===
current_prompt = """
Please read through the following witness statement and generate only the following two sections of a Circumstances Report:

1. Description of Events (Chronological)
- Present the events in the order they occurred.
- Use **third-person formal investigative tone**, as if writing a report **about the claimant**, not from their perspective.
- Begin each entry with the date in bold (e.g., **25 March 2025**).
- Use a new paragraph for each distinct event or observation, but do not begin that paragraph with a date.
- Do not include headings like “(continued)” or first-person language such as “I” or “my.”

2. Post-Incident Condition
- Present the post-incident effects based on what is stated in the statement.
- Maintain a professional tone as if writing a report, not repeating the witness’s words verbatim.
- Do not speculate or add fluff. Focus only on facts stated in the statement.

ONLY include factual information stated in the witness statement. Do not add opinions or assumptions.

Witness Statement:
{text}
""".strip()

# === PDF Text Extraction ===
def extract_text_from_pdf(file_path):
    doc = fitz.open(file_path)
    text = ''
    for page in doc:
        text += page.get_text()
    return text

# === Word Document Injection ===
def insert_paragraph_after(paragraph, text):
    parent = paragraph._parent
    new_para = parent.add_paragraph()

    # Bold the line if it starts with a date format like "25 March 2025"
    match = re.match(r"^(\d{1,2} \w+ \d{4})(.*)", text)
    if match:
        date_part = match.group(1)
        rest = match.group(2).strip()
        run = new_para.add_run(date_part)
        run.bold = True
        if rest:
            new_para.add_run(" " + rest)
    else:
        new_para.add_run(text)

    # Maintain proper ordering in the document
    p = new_para._element
    paragraph._element.addnext(p)
    return new_para

def insert_into_circumstances_section(template_path, output_path, ai_summary):
    doc = Document(template_path)

    found_heading = False
    summary_lines = ai_summary.strip().split("\n")

    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().upper() == "CIRCUMSTANCES":
            found_heading = True
            for j in range(i + 1, len(doc.paragraphs)):
                target = doc.paragraphs[j - 1]
                if target.text.strip():
                    print(f"Inserting after: {target.text}")
                    target.text = summary_lines[0].strip()
                    previous_para = target
                    for line in summary_lines[1:]:
                        new_para = insert_paragraph_after(previous_para, line.strip())
                        previous_para = new_para
                    break
            break

    if not found_heading:
        raise ValueError("Could not find 'CIRCUMSTANCES' section.")

    doc.save(output_path)

# === AI Summarization ===
def summarize_text(text, model, num_ctx=8192, temperature=0.0):
    filled_prompt = current_prompt.replace("{text}", text)
    response: ChatResponse = chat(
        model=model,
        messages=[{
            'role': 'user',
            'content': filled_prompt
        }],
        options={
            'num_ctx': num_ctx,
            'temperature': temperature
        }
    )
    return response.message.content

def clean_summary_text(text):
    # Expanded to remove common LLM preambles and wrap-up lines
    text = re.sub(r"(?i)^.*?(summary|report)\s*[:\-]?\s*", "", text, count=1)
    text = re.split(r"\n+(do you want me to.*|let me know if.*|please note.*|if you need more.*)", text, flags=re.IGNORECASE)[0]
    return text.strip()

def extract_title_from_summary(text):
    match = re.search(r'\*\*(.+?)\*\*', text)
    if match:
        title = match.group(1)
        title = re.sub(r'[\\/:"*?<>|]+', '', title).strip()
        return title
    return "summary_report"

# === PDF Summary Output ===
def write_summary_to_pdf(summary_text, output_path):
    doc = SimpleDocTemplate(output_path)
    styles = getSampleStyleSheet()

    paragraph_style = styles["BodyText"]
    bullet_style = ParagraphStyle(
        name="Bullet",
        parent=styles["Normal"],
        leftIndent=20,
        bulletIndent=10,
        spaceBefore=4,
    )
    heading_style = ParagraphStyle(
        name="Heading",
        parent=styles["Heading2"],
        spaceAfter=6,
    )

    elements = []

    def format_markdown(text):
        text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
        text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
        return text

    lines = summary_text.strip().split('\n')
    buffer = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if buffer:
                combined = format_markdown(' '.join(buffer))
                elements.append(Paragraph(combined, paragraph_style))
                elements.append(Spacer(1, 0.15 * inch))
                buffer = []
            continue

        if stripped.startswith("•"):
            bullet_content = stripped[1:].strip()
            bullet_formatted = format_markdown(bullet_content)
            elements.append(ListFlowable(
                [ListItem(Paragraph(bullet_formatted, bullet_style))],
                bulletType='bullet'
            ))
            continue

        if stripped.startswith("**") and stripped.endswith("**"):
            heading = stripped.strip("*")
            elements.append(Paragraph(f"<b>{heading}</b>", heading_style))
            continue

        buffer.append(stripped)

    if buffer:
        combined = format_markdown(' '.join(buffer))
        elements.append(Paragraph(combined, paragraph_style))

    doc.build(elements)

# === Pipeline ===
def process_pdf(pdf_path):
    try:
        pdf_text = extract_text_from_pdf(pdf_path)
        selected_model = model_var.get()

        start_time = time.time()
        temperature = 0.0
        summary = summarize_text(pdf_text, model=selected_model, num_ctx=num_ctx, temperature=temperature)
        end_time = time.time()
        elapsed_time = end_time - start_time
        print(f"[INFO] Model '{selected_model}' took {elapsed_time:.2f} seconds with temperature {temperature:.1f}.")

        cleaned_summary = clean_summary_text(summary)
        title = extract_title_from_summary(cleaned_summary)
        output_dir = os.path.dirname(pdf_path)
        output_file = os.path.join(output_dir, f"{title}.pdf")

        write_summary_to_pdf(cleaned_summary, output_file)

        template_path = "Proof - Report Template Blank.docx"
        word_output_path = os.path.join(output_dir, f"{title}.docx")
        insert_into_circumstances_section(template_path, word_output_path, cleaned_summary)

        messagebox.showinfo("Success", f"Summary saved to:\n{output_file}")
    except Exception as e:
        messagebox.showerror("Error", str(e))

# === GUI Callbacks ===
def select_file():
    file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
    if file_path:
        process_pdf(file_path)

def on_drop(event):
    pdf_path = event.data.strip('{}')
    if pdf_path.lower().endswith('.pdf'):
        process_pdf(pdf_path)
    else:
        messagebox.showwarning("Invalid File", "Please drop a .pdf file.")

# === GUI Setup ===
num_ctx = 8192

app = TkinterDnD.Tk()
app.title("PDF Summarizer")
app.geometry("550x350")

label = tk.Label(app, text="Drag and drop a PDF here,\nor click the button to select one.", font=("Helvetica", 14))
label.pack(pady=50)

drop_area = tk.Label(app, text="⬇ Drop PDF File Here ⬇", relief="groove", height=5, width=50)
drop_area.pack(pady=10)
drop_area.drop_target_register(DND_FILES)
drop_area.dnd_bind('<<Drop>>', on_drop)

browse_button = tk.Button(app, text="Select PDF File", command=select_file)
browse_button.pack(pady=20)

model_var = tk.StringVar(value="gemma3:4b")

model_label = tk.Label(app, text="Model:", font=("Helvetica", 10))
model_label.place(relx=0.71, rely=0.05)

model_menu = tk.OptionMenu(app, model_var, "gemma3:1b", "gemma3:4b", "gemma3:12b")
model_menu.place(relx=0.79, rely=0.04)

status_label = tk.Label(app, text="", font=("Helvetica", 10), fg="gray")
status_label.pack(pady=5)

def edit_prompt_window():
    def save_prompt():
        global current_prompt
        current_prompt = text_widget.get("1.0", tk.END).strip()
        prompt_editor.destroy()
        messagebox.showinfo("Prompt Updated", "Summarization prompt has been updated.")

    prompt_editor = tk.Toplevel(app)
    prompt_editor.title("Edit AI Prompt")
    prompt_editor.geometry("600x500")

    text_widget = tk.Text(prompt_editor, wrap=tk.WORD)
    text_widget.insert(tk.END, current_prompt)
    text_widget.pack(expand=True, fill="both", padx=10, pady=10)

    save_button = tk.Button(prompt_editor, text="Save Prompt", command=save_prompt)
    save_button.pack(pady=5)

edit_prompt_button = tk.Button(app, text="✎ Edit Prompt", command=edit_prompt_window)
edit_prompt_button.pack(pady=10)

app.mainloop()